# -*- coding: utf-8 -*-
"""
心理学报格式 Word 生成器
双模式：命令行 python gen_docx_dualcol.py 论文.md（AI Agent 用）
        双击 exe 弹出文件选择窗口（人手动用）
"""
import sys, io, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 核心生成逻辑
# ============================================================
def generate(input_file):
    """读取 .md 文件，生成 .docx，返回输出文件路径"""
    output_file = os.path.splitext(os.path.basename(input_file))[0] + '.docx'

    # ---- 辅助函数（必须在解析之前定义）----
    def normalize_punctuation(text):
        """将中文标点替换为英文标点：逗号/分号/冒号加空格，括号不加，日式引号转英文弯引号"""
        text = text.replace('，', ', ')
        text = text.replace('；', '; ')
        text = text.replace('：', ': ')
        text = text.replace('（', '(')
        text = text.replace('）', ')')
        text = text.replace('「', '“')   # → "（英文左弯引号，Times New Roman）
        text = text.replace('」', '”')   # → "（英文右弯引号，Times New Roman）
        return text

    def normalize_keep_colon(text):
        """替换括号、逗号、分号、引号，不动冒号（标题和参考文献共用）"""
        text = text.replace('，', ', ')
        text = text.replace('；', '; ')
        text = text.replace('（', '(')
        text = text.replace('）', ')')
        text = text.replace('「', '“')
        text = text.replace('」', '”')
        return text

    # ---- 读取并解析 Markdown ----
    with open(input_file, encoding='utf-8') as f:
        md_text = f.read()

    lines = md_text.split('\n')
    title_text = ""
    abstract_lines = []
    keywords_line = ""
    body_sections = []
    ref_lines = []
    current_level = None
    current_title = ""
    current_content = []
    in_abstract = False
    in_refs = False

    for line in lines:
        s = line.strip()
        if not s:
            if in_abstract and abstract_lines: in_abstract = False
            continue
        if s == '---': continue
        if s.startswith('# '): title_text = normalize_keep_colon(s[2:]); continue
        if s.startswith('**你的姓名**') or s.startswith('你的姓名'): continue
        if s == '## 摘要': in_abstract = True; continue
        if s.startswith('关键词'):
            keywords_line = normalize_punctuation(s.replace('**','').replace('关键词：','').replace('关键词','').replace('：','').strip())
            continue
        if s.startswith('## 参考文献'):
            in_refs = True
            if current_title:
                body_sections.append((current_level, current_title, current_content))
                current_content = []; current_title = ""; current_level = None
            continue
        if in_refs:
            if not s.startswith('*[') and not s.startswith('---'):
                ref_lines.append(normalize_punctuation(s))
            continue
        if in_abstract: abstract_lines.append(normalize_punctuation(s)); continue
        if s.startswith('## '):
            if current_title:
                body_sections.append((current_level, current_title, current_content))
            current_level = "h2"; current_title = normalize_keep_colon(s[3:]); current_content = []
            continue
        if s.startswith('### '):
            if current_title:
                body_sections.append((current_level, current_title, current_content))
            current_level = "h3"; current_title = normalize_keep_colon(s[4:]); current_content = []
            continue
        s = normalize_punctuation(s.replace('**', '').replace('*', ''))
        current_content.append(s)
    if current_title:
        body_sections.append((current_level, current_title, current_content))

    def add_mixed_run(p, text, cn_font='宋体', en_font='Times New Roman',
                      size_pt=10.5, bold=False, italic=False):
        runs_data = []
        current = ''
        current_type = None
        for ch in text:
            cp = ord(ch)
            is_cjk = (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
                      0xF900 <= cp <= 0xFAFF or 0x20000 <= cp <= 0x2A6DF or
                      ch in '。、？！【】《》—…αβ×')
            t = 'cn' if is_cjk else 'en'
            if t == current_type: current += ch
            else:
                if current: runs_data.append((current, current_type))
                current = ch; current_type = t
        if current: runs_data.append((current, current_type))
        for seg_text, seg_type in runs_data:
            run = p.add_run(seg_text)
            run.font.size = Pt(size_pt)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), cn_font)
            run.font.name = cn_font if seg_type == 'cn' else en_font
            rFonts.set(qn('w:ascii'), en_font)
            rFonts.set(qn('w:hAnsi'), en_font)
            for tag, flag in [('w:b', bold), ('w:i', italic)]:
                existing = rpr.find(qn(tag))
                if existing is not None: rpr.remove(existing)
                if flag: rpr.append(OxmlElement(tag))

    def add_text_with_italics(p, text, cn_font='宋体', en_font='Times New Roman',
                              size_pt=10.5, bold=False, base_italic=False):
        """解析 ^...^ 斜体标记并渲染。^ 标记内容强制斜体，base_italic 用于参考文献 * 斜体叠加。"""
        parts = text.split('^')
        for i, part in enumerate(parts):
            if not part:
                continue
            is_italic = (i % 2 == 1) or base_italic
            add_mixed_run(p, part, cn_font=cn_font, en_font=en_font,
                          size_pt=size_pt, bold=bold, italic=is_italic)

    def set_multiline(p, ratio=1.16):
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = ratio

    def add_ref_paragraph(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_multiline(p, 1.16)
        parts = text.split('*')
        for idx, part in enumerate(parts):
            if not part: continue
            add_text_with_italics(p, part, size_pt=9, base_italic=(idx % 2 == 1))
        return p

    def set_page(section):
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.00)
        section.right_margin = Cm(2.00)

    # ---- 构建文档 ----
    doc = Document()
    sec0 = doc.sections[0]
    set_page(sec0)

    # 题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_multiline(p, 1.16)
    add_text_with_italics(p, title_text, cn_font='黑体', size_pt=22)

    # 作者
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_multiline(p, 1.16)
    add_mixed_run(p, '你的姓名', cn_font='仿宋', size_pt=14)

    # 单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_multiline(p, 1.16)
    add_mixed_run(p, '(你的学校 你的学院, 城市 邮编)', cn_font='宋体', size_pt=9)

    # 摘要
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.right_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(8)
    set_multiline(p, 1.16)
    add_mixed_run(p, '摘  要  ', size_pt=10.5, bold=True)
    add_text_with_italics(p, abstract_lines[0] if abstract_lines else '', size_pt=10.5)

    # 关键词
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.right_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(24)
    set_multiline(p, 1.16)
    add_mixed_run(p, '关键词  ', size_pt=10.5, bold=True)
    add_text_with_italics(p, keywords_line, size_pt=10.5)

    # ---- 双栏正文 ----
    sec1 = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    set_page(sec1)
    existing_cols = sec1._sectPr.findall(qn('w:cols'))
    for ec in existing_cols:
        sec1._sectPr.remove(ec)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '435')
    cols.set(qn('w:equalWidth'), '1')
    sec1._sectPr.append(cols)

    h1_num = 0; h2_num = 0
    for level, title, content_lines in body_sections:
        if level == "h2":
            h1_num += 1; h2_num = 0
            desc = title.split(" ", 1)[-1] if " " in title else title
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Cm(0)
            set_multiline(p, 1.16)
            add_text_with_italics(p, str(h1_num) + '  ' + desc, cn_font='宋体', size_pt=14)
        else:
            h2_num += 1
            desc = title.split(" ", 1)[-1] if " " in title else title
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(0)
            set_multiline(p, 1.16)
            add_text_with_italics(p, str(h1_num) + '.' + str(h2_num) + '  ' + desc, cn_font='黑体', size_pt=10.5)
        for cline in content_lines:
            cline = cline.strip()
            if not cline: continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_multiline(p, 1.16)
            add_text_with_italics(p, cline, size_pt=10.5)

    # ---- 参考文献排序 ----
    try:
        from pypinyin import pinyin, Style
        _HAS_PYPINYIN = True
    except ImportError:
        _HAS_PYPINYIN = False

    def sort_key_ref(line):
        first_char = line.strip()[0] if line.strip() else ''
        if '一' <= first_char <= '鿿':
            surname = ''
            for ch in line.strip():
                if '一' <= ch <= '鿿':
                    surname += ch
                else:
                    break
            if _HAS_PYPINYIN:
                py = pinyin(surname, style=Style.TONE3)
                py_str = ''.join([p[0] for p in py])
                return ('0_' + py_str + '_' + surname)
            else:
                return ('0_' + surname)
        else:
            return ('1_' + line.strip().lower())

    ref_lines.sort(key=sort_key_ref)

    # 参考文献标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    set_multiline(p, 1.16)
    add_mixed_run(p, '参 考 文 献', cn_font='黑体', size_pt=12, bold=True)
    for r_line in ref_lines:
        r_line = r_line.strip()
        if not r_line: continue
        add_ref_paragraph(r_line)

    doc.save(output_file)
    return output_file


# ============================================================
# 入口：CLI 模式（有命令行参数）或 GUI 模式（双击）
# ============================================================
if __name__ == '__main__':
    if len(sys.argv) > 1:
        # CLI 模式：AI Agent / 命令行使用
        input_file = sys.argv[1]
        print('Input : ' + input_file)
        output = generate(input_file)
        print('DONE: ' + output)
        print('Size: ' + str(os.path.getsize(output)))
    else:
        # GUI 模式：用户双击启动
        import tkinter as tk
        from tkinter import filedialog, messagebox

        def select_file():
            path = filedialog.askopenfilename(
                title='选择 Markdown 论文文件',
                filetypes=[('Markdown 文件', '*.md'), ('所有文件', '*.*')]
            )
            if path:
                file_var.set(path)

        def do_generate():
            path = file_var.get().strip()
            if not path:
                messagebox.showwarning('提示', '请先选择一个 .md 文件')
                return
            if not os.path.exists(path):
                messagebox.showerror('错误', '文件不存在：\n' + path)
                return
            try:
                status_var.set('正在生成，请稍候……')
                root.update()
                out = generate(path)
                status_var.set('生成成功！→ ' + os.path.abspath(out))
            except Exception as e:
                status_var.set('生成失败：' + str(e))
                messagebox.showerror('错误', '生成过程中出错：\n' + str(e))

        root = tk.Tk()
        root.title('心理学报格式生成器')
        root.resizable(False, False)

        # 窗口居中
        w, h = 520, 170
        x = (root.winfo_screenwidth() - w) // 2
        y = (root.winfo_screenheight() - h) // 2
        root.geometry(f'{w}x{h}+{x}+{y}')

        frame = tk.Frame(root, padx=16, pady=16)
        frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(frame, text='选择 Markdown 论文文件：', font=('Microsoft YaHei UI', 10)).pack(anchor=tk.W)

        row = tk.Frame(frame)
        row.pack(fill=tk.X, pady=(6, 10))
        file_var = tk.StringVar()
        tk.Entry(row, textvariable=file_var, font=('Microsoft YaHei UI', 9)).pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Button(row, text='浏览…', command=select_file, width=8,
                  font=('Microsoft YaHei UI', 9)).pack(side=tk.LEFT, padx=(6, 0))

        tk.Button(frame, text='生成 Word 文档', command=do_generate,
                  font=('Microsoft YaHei UI', 10, 'bold'), height=2).pack(fill=tk.X)

        status_var = tk.StringVar(value='就绪 — 请选择 .md 文件后点击生成')
        tk.Label(frame, textvariable=status_var, fg='#555',
                 font=('Microsoft YaHei UI', 9)).pack(anchor=tk.W, pady=(8, 0))

        root.mainloop()
